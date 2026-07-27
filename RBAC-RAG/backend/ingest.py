"""File parsing + chunking utilities.

- Parse TXT/MD/PDF/DOCX -> plain text
- Chunk into ~500-token windows with 50-token overlap using tiktoken
"""
import io
import re
import logging
from typing import Iterable
import tiktoken

logger = logging.getLogger(__name__)

_ENC = tiktoken.get_encoding("cl100k_base")


def parse_txt(data: bytes) -> str:
    for enc in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def parse_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    pages = []
    for p in reader.pages:
        try:
            pages.append(p.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages)


def parse_docx(data: bytes) -> str:
    import docx  # python-docx
    doc = docx.Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def parse_file(filename: str, data: bytes) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(data)
    if lower.endswith(".docx"):
        return parse_docx(data)
    if lower.endswith(".txt") or lower.endswith(".md") or lower.endswith(".markdown"):
        return parse_txt(data)
    raise ValueError(f"Unsupported file type: {filename}")


def _clean_text(text: str) -> str:
    # Collapse whitespace and strip control chars
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    text = re.sub(r"[ \t]+\n", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def chunk_text(text: str, chunk_tokens: int = 500, overlap_tokens: int = 50) -> list[str]:
    """Token-based sliding window chunker."""
    text = _clean_text(text)
    if not text:
        return []
    tokens = _ENC.encode(text)
    if len(tokens) <= chunk_tokens:
        return [text]
    chunks: list[str] = []
    step = chunk_tokens - overlap_tokens
    if step <= 0:
        step = chunk_tokens
    for start in range(0, len(tokens), step):
        window = tokens[start : start + chunk_tokens]
        if not window:
            break
        chunk = _ENC.decode(window).strip()
        if chunk:
            chunks.append(chunk)
        if start + chunk_tokens >= len(tokens):
            break
    return chunks
